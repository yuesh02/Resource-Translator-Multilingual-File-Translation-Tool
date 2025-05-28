import os
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from PIL import Image
import pytesseract
import fitz
from docx import Document
from pdf2docx import Converter
from googletrans import Translator

# Define supported target languages
target_languages = {
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "Hindi": "hi",
    "Marathi": "mr",
    "Bengali": "bn",
    "Gujarati": "gu",
    "Tamil": "ta",
    "Telugu": "te",
    # Add more languages as needed
}

# Translation functions
def translate_image(image, target_language):
    return pytesseract.image_to_string(image, lang=target_language)

def translate_pdf_document(file_path, target_language):
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

def translate_word_document(file_path, target_language):
    doc = Document(file_path)
    translated_doc = Document()
    for paragraph in doc.paragraphs:
        translated_text = translate_text(paragraph.text, target_language)
        translated_doc.add_paragraph(translated_text)
    translated_file_path = os.path.splitext(file_path)[0] + "_translated.docx"
    translated_doc.save(translated_file_path)
    return translated_file_path

def translate_text(text, target_language):
    translator = Translator()
    translated_text = translator.translate(text, dest=target_language)
    return translated_text.text

def translate_text_file(file_path, target_language):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    translated_text = translate_text(text, target_language)
    return translated_text

# Function for downloading files
def download_file(content, file_extension=".txt"):
    file_path = filedialog.asksaveasfilename(defaultextension=file_extension)
    if file_path:
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(content)
        messagebox.showinfo("Download Complete", "File downloaded successfully.")

# Main translation logic
def translate_resource():
    file_path = entry_file_path.get()
    if not file_path:
        messagebox.showerror("Error", "Please select a file.")
        return

    target_language = target_languages.get(combobox_language.get())
    if not target_language:
        messagebox.showerror("Error", "Please select a target language.")
        return

    resource_type = combobox_resource_type.get()
    if resource_type == "Image":
        image = Image.open(file_path)
        translated_text = translate_image(image, target_language)
    elif resource_type == "PDF Document":
        translated_text = translate_pdf_document(file_path, target_language)
    elif resource_type == "Word Document":
        translated_file_path = translate_word_document(file_path, target_language)
        messagebox.showinfo("Translation Complete", f"Translation saved to {translated_file_path}")
        return
    elif resource_type == "Text File":
        translated_text = translate_text_file(file_path, target_language)
    else:
        messagebox.showerror("Error", "Invalid resource type.")
        return

    download_file(translated_text)

# Function for opening files
def open_file():
    file_path = filedialog.askopenfilename()
    entry_file_path.delete(0, tk.END)
    entry_file_path.insert(0, file_path)

# Create GUI window
window = tk.Tk()
window.title("Resource Translator")

# Widgets
label_file_path = ttk.Label(window, text="File Path:")
label_file_path.grid(row=0, column=0, padx=5, pady=5, sticky="e")

entry_file_path = ttk.Entry(window, width=50)
entry_file_path.grid(row=0, column=1, columnspan=2, padx=5, pady=5)

button_open_file = ttk.Button(window, text="Open File", command=open_file)
button_open_file.grid(row=0, column=3, padx=5, pady=5)

label_language = ttk.Label(window, text="Target Language:")
label_language.grid(row=1, column=0, padx=5, pady=5, sticky="e")

combobox_language = ttk.Combobox(window, values=list(target_languages.keys()), state="readonly")
combobox_language.grid(row=1, column=1, padx=5, pady=5)

label_resource_type = ttk.Label(window, text="Resource Type:")
label_resource_type.grid(row=1, column=2, padx=5, pady=5, sticky="e")

combobox_resource_type = ttk.Combobox(window, values=["Image", "PDF Document", "Word Document", "Text File"], state="readonly")
combobox_resource_type.grid(row=1, column=3, padx=5, pady=5)

button_translate = ttk.Button(window, text="Translate", command=translate_resource)
button_translate.grid(row=2, column=0, columnspan=4, padx=5, pady=5)

# Run the GUI application
window.mainloop()


